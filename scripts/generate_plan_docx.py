"""Generate PRO HR Implementation Plan as a Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Style setup ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_colored_heading(text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def make_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return t

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

# ════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("PRO HR")
r.font.size = Pt(42)
r.font.color.rgb = RGBColor(0x3b, 0x82, 0xf6)
r.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Implementation Plan for Platform Updates")
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Version 2.0  •  May 2026\nPrepared for: Astral Cassini Development Team")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_paragraph()
goal = doc.add_paragraph()
goal.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = goal.add_run("Goal: Elevate composite quality score from 71/100 (B) → 90+/100 (A)")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)
r.bold = True

doc.add_page_break()

# ════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════
add_colored_heading("Table of Contents", 1, RGBColor(0x1a, 0x1a, 0x2e))
toc_items = [
    "1. Executive Summary",
    "2. Current Score Breakdown",
    "3. Phase 1 — Critical Fixes (6 Tasks)",
    "4. Phase 2 — Reliability Improvements (5 Tasks)",
    "5. Phase 3 — Polish & Documentation (4 Tasks)",
    "6. Effort Summary & Timeline",
    "7. Verification Checklist",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════
add_colored_heading("1. Executive Summary", 1, RGBColor(0x1a, 0x1a, 0x2e))

doc.add_paragraph(
    "PRO HR is a multi-agent recruitment orchestration platform featuring a deterministic "
    "state machine backend, FAISS-backed RAG pipeline, real-time WebSocket streaming, and a "
    "glassmorphic Next.js frontend. A comprehensive deep audit identified 30+ issues across "
    "9 quality dimensions."
)
doc.add_paragraph(
    "This document presents a structured, phased implementation plan with 15 tasks organized "
    "into 3 phases — Critical Fixes, Reliability Improvements, and Polish — to systematically "
    "resolve every identified gap and bring the platform to production-grade quality."
)

p = doc.add_paragraph()
r = p.add_run("Key Findings:")
r.bold = True
r.font.size = Pt(12)

add_bullet("Dashboard 'Candidates in Pipeline' metric permanently displays 0 (broken backend query)", "Bug: ")
add_bullet("Login page branding says 'AGENTHIRE' while entire app uses 'PRO HR'", "UX: ")
add_bullet("Unused TailwindCSS import adds ~200KB to CSS bundle", "Performance: ")
add_bullet("No React Error Boundaries — component errors crash to white screen", "Stability: ")
add_bullet("No rate limiting on login endpoint — brute-force vulnerable", "Security: ")
add_bullet("Dead graph/ module with empty __init__.py", "Cleanup: ")

doc.add_page_break()

# ════════════════════════════════════════════════
# 2. CURRENT SCORE BREAKDOWN
# ════════════════════════════════════════════════
add_colored_heading("2. Current Score Breakdown", 1, RGBColor(0x1a, 0x1a, 0x2e))

doc.add_paragraph("The following scores were determined through code review, runtime testing, and UI inspection:")

make_table(
    ["Dimension", "Score", "Grade", "Key Issue"],
    [
        ["Architecture & Design", "82/100", "A-", "Workflow state duplication in JSON blob"],
        ["Backend Quality", "78/100", "B+", "No rate limiting on auth endpoints"],
        ["Frontend Quality", "73/100", "B", "Massive inline styles, unused Tailwind"],
        ["Test Coverage", "71/100", "B", "No WebSocket or RAG tests"],
        ["Security", "74/100", "B", "Missing CSP headers"],
        ["UX / Visual Design", "80/100", "A-", "Empty states lack guidance"],
        ["DevOps & CI/CD", "65/100", "C+", "CI workflows not functional"],
        ["Documentation", "60/100", "C", "No architecture or API guide"],
        ["Production Readiness", "58/100", "C", "SQLite only, no health checks"],
        ["COMPOSITE", "71/100", "B", "—"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════
# 3. PHASE 1
# ════════════════════════════════════════════════
add_colored_heading("3. Phase 1 — Critical Fixes", 1, RGBColor(0xf4, 0x3f, 0x5e))
doc.add_paragraph("These tasks fix broken functionality, security gaps, and dead code. They are independent, low-risk, and should be completed first.")

# Task 1.1
add_colored_heading('Task 1.1 — Fix "Candidates in Pipeline" Metric', 2, RGBColor(0x3b, 0x82, 0xf6))
make_table(["Attribute", "Detail"], [
    ["Priority", "🔴 P0 — Visibly broken feature"],
    ["Effort", "15 minutes"],
    ["Files", "backend/app/core/orchestrator.py (line 462-469)"],
    ["Impact", "Dashboard shows correct candidate count"],
])
doc.add_paragraph()
doc.add_paragraph("Root Cause: The _job_to_summary() function does not include candidates_count in its return dictionary. The frontend's totalCandidates memo sums j.candidates_count which is always undefined, resulting in 0.")
p = doc.add_paragraph()
r = p.add_run("Fix: ")
r.bold = True
p.add_run("Parse the workflow_state JSON blob inside _job_to_summary() and add candidates_count = len(state.get('candidates', [])) to the returned dict.")
doc.add_paragraph("Verification: Dashboard shows non-zero 'Candidates in Pipeline' when jobs contain candidates.")

# Task 1.2
add_colored_heading("Task 1.2 — Unify Branding (PRO HR Everywhere)", 2, RGBColor(0x3b, 0x82, 0xf6))
make_table(["Attribute", "Detail"], [
    ["Priority", "🟡 P1 — Brand confusion"],
    ["Effort", "5 minutes"],
    ["Files", "frontend/src/app/login/page.tsx (line 134-139)"],
    ["Impact", "Consistent branding across entire application"],
])
doc.add_paragraph()
doc.add_paragraph('Fix: Replace "AGENTHIRE" with "PRO HR" and update the subtitle from "Multi-Agent Recruitment Ecosystem" to "Autonomous Recruitment Ecosystem" to match the landing page.')
doc.add_paragraph("Verification: Navigate to /login — brand name matches sidebar, landing page, and document title.")

# Task 1.3
add_colored_heading("Task 1.3 — Remove Unused Tailwind Import", 2, RGBColor(0x3b, 0x82, 0xf6))
make_table(["Attribute", "Detail"], [
    ["Priority", "🟡 P1 — ~200KB unused CSS in bundle"],
    ["Effort", "5 minutes"],
    ["Files", "frontend/src/app/globals.css (line 1), frontend/package.json"],
    ["Impact", "Significantly smaller CSS bundle"],
])
doc.add_paragraph()
doc.add_paragraph('Fix: Remove @import "tailwindcss" from line 1 of globals.css. Remove @tailwindcss/postcss and tailwindcss from devDependencies in package.json. Update postcss.config.mjs if it references the tailwind plugin.')
doc.add_paragraph("Verification: npm run build succeeds. CSS bundle size drops significantly.")

# Task 1.4
add_colored_heading("Task 1.4 — Add React Error Boundary", 2, RGBColor(0x3b, 0x82, 0xf6))
make_table(["Attribute", "Detail"], [
    ["Priority", "🟡 P1 — White screen crashes"],
    ["Effort", "15 minutes"],
    ["Files", "NEW: frontend/src/components/ErrorBoundary.tsx, EDIT: frontend/src/app/layout.tsx"],
    ["Impact", "Graceful error handling instead of blank screen"],
])
doc.add_paragraph()
doc.add_paragraph("Fix: Create a class-based ErrorBoundary component with getDerivedStateFromError and componentDidCatch. It renders a styled error card with a 'Reload Page' button. Wrap <MainContent>{children}</MainContent> in layout.tsx with <ErrorBoundary>.")
doc.add_paragraph("Verification: Throw a test error in any component — error card appears instead of white screen.")

# Task 1.5
add_colored_heading("Task 1.5 — Add Login Rate Limiting", 2, RGBColor(0x3b, 0x82, 0xf6))
make_table(["Attribute", "Detail"], [
    ["Priority", "🟡 P1 — Brute-force vulnerability"],
    ["Effort", "20 minutes"],
    ["Files", "backend/app/main.py, backend/app/api/auth.py, requirements.txt"],
    ["Impact", "Login endpoint limited to 5 attempts per minute per IP"],
])
doc.add_paragraph()
doc.add_paragraph("Fix: Install slowapi. Initialize a Limiter instance in main.py with get_remote_address as the key function. Apply @limiter.limit('5/minute') decorator to the /api/auth/login endpoint. Add the RateLimitExceeded exception handler to the FastAPI app.")
doc.add_paragraph("Verification: After 5 rapid login attempts, receive HTTP 429 Too Many Requests.")

# Task 1.6
add_colored_heading("Task 1.6 — Delete Dead graph/ Module", 2, RGBColor(0x3b, 0x82, 0xf6))
make_table(["Attribute", "Detail"], [
    ["Priority", "🟢 P2 — Dead code"],
    ["Effort", "2 minutes"],
    ["Files", "DELETE: backend/app/graph/ (entire directory)"],
    ["Impact", "Cleaner codebase, less confusion for contributors"],
])
doc.add_paragraph()
doc.add_paragraph("Fix: Delete the backend/app/graph/ directory which contains only an empty __init__.py. This module was part of an earlier LangGraph implementation that has been fully replaced by the deterministic Orchestrator class.")
doc.add_paragraph("Verification: pytest tests/ -v still passes (47/47 tests).")

doc.add_page_break()

# ════════════════════════════════════════════════
# 4. PHASE 2
# ════════════════════════════════════════════════
add_colored_heading("4. Phase 2 — Reliability Improvements", 1, RGBColor(0xf5, 0x9e, 0x0b))
doc.add_paragraph("These tasks improve security posture, mobile UX, test coverage, and CI automation. They build on Phase 1 foundations.")

# Task 2.1
add_colored_heading("Task 2.1 — Add Content-Security-Policy Headers", 2, RGBColor(0x3b, 0x82, 0xf6))
make_table(["Attribute", "Detail"], [
    ["Priority", "🟡 P1 — XSS mitigation"],
    ["Effort", "15 minutes"],
    ["Files", "frontend/next.config.ts"],
    ["Impact", "Browser-enforced XSS protection"],
])
doc.add_paragraph()
doc.add_paragraph("Fix: Add async headers() function to next.config.ts returning Content-Security-Policy, X-Frame-Options (DENY), X-Content-Type-Options (nosniff), and Referrer-Policy headers for all routes.")

# Task 2.2
add_colored_heading("Task 2.2 — Fix Mobile Responsiveness", 2, RGBColor(0x3b, 0x82, 0xf6))
make_table(["Attribute", "Detail"], [
    ["Priority", "🟡 P1 — Broken on tablets/phones"],
    ["Effort", "25 minutes"],
    ["Files", "frontend/src/app/globals.css, jobs/page.tsx, layout.tsx, SidebarAndHeader.tsx"],
    ["Impact", "Usable layout on screens < 1024px"],
])
doc.add_paragraph()
doc.add_paragraph("Fix: Add @media (max-width: 1024px) breakpoints to collapse the sidebar, convert the Jobs page two-column grid to single-column, and remove the fixed marginLeft on mobile. Add a hamburger toggle for the sidebar.")

# Task 2.3
add_colored_heading("Task 2.3 — Improve Empty State UX", 2, RGBColor(0x3b, 0x82, 0xf6))
make_table(["Attribute", "Detail"], [
    ["Priority", "🟡 P1 — Poor first impression"],
    ["Effort", "30 minutes"],
    ["Files", "NEW: frontend/src/components/EmptyState.tsx, EDIT: candidates, interviews, audit pages"],
    ["Impact", "Guided onboarding for new users"],
])
doc.add_paragraph()
doc.add_paragraph("Fix: Create a reusable EmptyState component with icon, title, description, and optional CTA button. Replace bare 'Select a pipeline...' dropdowns in Candidates, Interviews, and Audit Trail pages with rich empty states that guide users to create their first job pipeline.")

# Task 2.4
add_colored_heading("Task 2.4 — Add WebSocket & RAG Test Coverage", 2, RGBColor(0x3b, 0x82, 0xf6))
make_table(["Attribute", "Detail"], [
    ["Priority", "🟡 P1 — Core features untested"],
    ["Effort", "45 minutes"],
    ["Files", "NEW: backend/tests/unit/test_websocket.py, backend/tests/unit/test_embeddings.py"],
    ["Impact", "Regression safety for real-time and search features"],
])
doc.add_paragraph()
doc.add_paragraph("Fix: Write WebSocket tests covering: connection rejection without token, ticket/job mismatch rejection, successful connection and heartbeat. Write embedding tests covering: index_resume adds documents, search_resumes returns ranked results, mock fallback behavior.")

# Task 2.5
add_colored_heading("Task 2.5 — Add GitHub Actions CI Pipeline", 2, RGBColor(0x3b, 0x82, 0xf6))
make_table(["Attribute", "Detail"], [
    ["Priority", "🟡 P1 — No automated quality gates"],
    ["Effort", "30 minutes"],
    ["Files", ".github/workflows/ci.yml"],
    ["Impact", "Automated test + lint + build on every push/PR"],
])
doc.add_paragraph()
doc.add_paragraph("Fix: Update ci.yml with two jobs — backend (Python 3.14, pytest) and frontend (Node 22, lint + test + build). Triggered on push to master/main and all PRs. Uses environment variables for SECRET_KEY and DATABASE_URL so tests run without .env files.")

doc.add_page_break()

# ════════════════════════════════════════════════
# 5. PHASE 3
# ════════════════════════════════════════════════
add_colored_heading("5. Phase 3 — Polish & Documentation", 1, RGBColor(0x10, 0xb9, 0x81))
doc.add_paragraph("These tasks improve long-term maintainability and developer experience. They can be implemented incrementally.")

tasks_p3 = [
    ("Task 3.1 — Extract Inline Styles to CSS Modules", "🟢 P2", "2-3 hours", 
     "All page components (landing, login, sidebar)",
     "Migrate inline style={{}} objects to .module.css files with semantic class names. Start with landing page (450 lines, 90% inline) and login page (294 lines). Replace onMouseOver/onMouseOut JS handlers with CSS :hover rules."),
    ("Task 3.2 — Add Architecture Documentation", "🟢 P2", "30 minutes",
     "NEW: docs/architecture.md",
     "Create architecture.md with system diagram, 7-agent pipeline flow with stage names, database schema overview, WebSocket event types, and HITL breakpoint rules."),
    ("Task 3.3 — Add API Integration Guide", "🟢 P2", "30 minutes",
     "NEW: docs/api-guide.md",
     "Document auth flow (login → cookie → CSRF), WebSocket handshake, job lifecycle API calls, resume upload constraints, and example cURL commands."),
    ("Task 3.4 — Add Frontend Negative-Path Tests", "🟢 P2", "45 minutes",
     "NEW: frontend/e2e/auth-negative.spec.ts",
     "Test invalid login credentials, expired session redirect, API server down toast, XSS payload sanitization, and unauthenticated dashboard access redirect."),
]
for title, pri, effort, files, desc in tasks_p3:
    add_colored_heading(title, 2, RGBColor(0x3b, 0x82, 0xf6))
    make_table(["Attribute", "Detail"], [["Priority", pri], ["Effort", effort], ["Files", files]])
    doc.add_paragraph()
    doc.add_paragraph(desc)

doc.add_page_break()

# ════════════════════════════════════════════════
# 6. EFFORT SUMMARY
# ════════════════════════════════════════════════
add_colored_heading("6. Effort Summary & Timeline", 1, RGBColor(0x1a, 0x1a, 0x2e))

make_table(
    ["Phase", "Tasks", "Total Effort", "Score Impact"],
    [
        ["Phase 1 — Critical", "6 tasks", "~1 hour", "71 → 80"],
        ["Phase 2 — Reliability", "5 tasks", "~2.5 hours", "80 → 88"],
        ["Phase 3 — Polish", "4 tasks", "~4 hours", "88 → 92+"],
        ["TOTAL", "15 tasks", "~7.5 hours", "71 → 92+"],
    ]
)

doc.add_paragraph()
doc.add_paragraph("All Phase 1 tasks are independent and can be executed in parallel by different team members. Phase 2 tasks have light dependencies on Phase 1 completion. Phase 3 tasks are fully incremental.")

doc.add_page_break()

# ════════════════════════════════════════════════
# 7. VERIFICATION CHECKLIST
# ════════════════════════════════════════════════
add_colored_heading("7. Verification Checklist", 1, RGBColor(0x1a, 0x1a, 0x2e))
doc.add_paragraph("After all changes are implemented, verify the following:")

checks = [
    'python -m pytest tests/ -v — 47+ tests passing',
    'npm run lint — 0 errors',
    'npm test -- --ci — 5+ tests passing',
    'npm run build — Clean build (15/15 pages)',
    'Dashboard "Candidates in Pipeline" shows real count',
    'Login page says "PRO HR" (not AGENTHIRE)',
    'CSS bundle ≤ 15KB (down from ~200KB)',
    'Error boundary catches component errors gracefully',
    '/api/auth/login returns 429 after 5 rapid attempts',
    'No backend/app/graph/ directory exists',
    'Mobile layout at 768px — no horizontal overflow',
    'GitHub Actions CI runs green on push',
    'Response headers include Content-Security-Policy',
    'Empty states show guidance and CTA buttons',
]
for c in checks:
    add_bullet(c)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Document prepared by: ")
r.bold = True
p.add_run("Antigravity AI Analysis Engine")

# ── Save ──
out_path = os.path.join(os.path.dirname(__file__), "..", "PRO_HR_Implementation_Plan.docx")
out_path = os.path.abspath(out_path)
doc.save(out_path)
print(f"✅ Document saved to: {out_path}")
