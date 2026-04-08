"""One-off generator: PRO HR clone & run guide as Word (.docx). Run from repo root: python scripts/build_clone_guide_docx.py"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    out = root / "docs" / "PRO-HR-Clone-and-Run-Guide.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("PRO HR (astral-cassini)", 0)
    p = doc.add_paragraph()
    p.add_run("Clone and run from GitHub").bold = True
    doc.add_paragraph(
        "This guide explains how to obtain the project from Git and run the FastAPI backend "
        "and Next.js frontend on your machine."
    )

    doc.add_heading("Prerequisites", level=1)
    for line in (
        "Git",
        "Node.js 20 or later and npm",
        "Python 3.11.x (recommended; matches CI and Docker)",
        'Optional: OpenAI API key for full agent and embedding behavior',
    ):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("1. Clone the repository", level=1)
    doc.add_paragraph("Open a terminal and run:", style=None)
    doc.add_paragraph(
        "git clone https://github.com/raviprateek2-dotcom/astral-cassini.git\n"
        "cd astral-cassini",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "If you use SSH: git clone git@github.com:raviprateek2-dotcom/astral-cassini.git"
    )

    doc.add_heading("2. Backend setup", level=1)
    doc.add_paragraph("From the repository root:", style=None)
    doc.add_paragraph("cd backend", style="Intense Quote")
    doc.add_paragraph("Create and activate a virtual environment (recommended):", style=None)
    doc.add_paragraph("Windows (PowerShell):", style=None)
    doc.add_paragraph(
        "python -m venv .venv\n"
        ".\\.venv\\Scripts\\Activate.ps1\n"
        "pip install -r requirements.txt",
        style="Intense Quote",
    )
    doc.add_paragraph("macOS / Linux:", style=None)
    doc.add_paragraph(
        "python3 -m venv .venv\n"
        "source .venv/bin/activate\n"
        "pip install -r requirements.txt",
        style="Intense Quote",
    )

    doc.add_heading("2a. Environment file", level=2)
    doc.add_paragraph(
        "The file backend/.env is not stored in Git. Copy the example file and edit it:"
    )
    doc.add_paragraph("Windows: copy .env.example .env", style="Intense Quote")
    doc.add_paragraph("macOS / Linux: cp .env.example .env", style="Intense Quote")
    doc.add_paragraph("You must set at least:", style=None)
    for line in (
        "SECRET_KEY — a random string with at least 32 characters (required for JWT and startup).",
        "OPENAI_API_KEY — if you use LLM and RAG features.",
        "DATABASE_URL — SQLite default in .env.example is fine for local development.",
    ):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("2b. Run the API", level=2)
    doc.add_paragraph(
        "With the virtual environment active and your working directory still backend/:",
        style=None,
    )
    doc.add_paragraph(
        "uvicorn app.main:app --reload --host 127.0.0.1 --port 8000",
        style="Intense Quote",
    )
    doc.add_paragraph("API base URL: http://127.0.0.1:8000")
    doc.add_paragraph("Interactive docs (Swagger): http://127.0.0.1:8000/docs")
    doc.add_paragraph("Keep this terminal open while you develop.")

    doc.add_heading("3. Frontend setup", level=1)
    doc.add_paragraph("Open a second terminal. From the repository root:", style=None)
    doc.add_paragraph("cd frontend", style="Intense Quote")
    doc.add_paragraph("npm install", style="Intense Quote")
    doc.add_paragraph(
        "Optional: if frontend/.env.example exists, copy it to .env.local and set BACKEND_URL "
        "if your API is not at http://127.0.0.1:8000. By default, Next.js rewrites /api to that host."
    )
    doc.add_paragraph("Start the development server:", style=None)
    doc.add_paragraph("npm run dev", style="Intense Quote")
    doc.add_paragraph(
        "Open the app in a browser at http://localhost:3000 (include the port; http://localhost alone is not enough)."
    )

    doc.add_heading("4. First login", level=1)
    doc.add_paragraph(
        "Register a user through POST /api/auth/register using the Swagger UI at /docs, "
        "then sign in at http://localhost:3000/login."
    )
    doc.add_paragraph(
        "For development only, you can enable seeded demo users in backend/.env "
        "(see .env.example: SEED_DEMO_USERS, DEMO_ADMIN_PASSWORD, DEMO_HR_PASSWORD). "
        "Passwords must be at least 8 characters."
    )

    doc.add_heading("5. Optional: automated verification", level=1)
    doc.add_paragraph(
        "From the repository root, on Linux/macOS (or Git Bash on Windows) you can run:"
    )
    doc.add_paragraph("bash scripts/verify-all.sh", style="Intense Quote")
    doc.add_paragraph(
        "On Windows, PowerShell equivalents live under scripts/ (for example verify-all.ps1). "
        "CI uses Python 3.11; use the same if local pip install fails on newer Python versions."
    )

    doc.add_heading("Quick reference", level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Step", "Action"),
        ("Clone", "git clone … then cd astral-cassini"),
        ("Backend deps", "cd backend → venv → pip install -r requirements.txt"),
        ("Backend config", "copy .env.example to .env; set SECRET_KEY"),
        ("Backend run", "uvicorn app.main:app --reload"),
        ("Frontend", "cd frontend → npm install → npm run dev"),
    ]
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b

    doc.add_paragraph("")
    doc.add_paragraph(
        "Repository: https://github.com/raviprateek2-dotcom/astral-cassini"
    ).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
