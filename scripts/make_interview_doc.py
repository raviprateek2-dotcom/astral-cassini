"""
Generates INTERVIEW_PREP.docx — a professionally formatted Word document
containing the complete PRO HR project interview preparation guide.
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_BG   = RGBColor(0x0F, 0x17, 0x2A)   # navy
ACCENT    = RGBColor(0x6C, 0x63, 0xFF)   # violet
GOLD      = RGBColor(0xFF, 0xC1, 0x07)   # amber
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_TXT = RGBColor(0x37, 0x41, 0x51)   # slate-700
DIVIDER   = RGBColor(0xE5, 0xE7, 0xEB)   # gray-200


def shade_paragraph(para, hex_fill: str):
    """Apply a solid background shading to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    pPr.append(shd)


def add_run(para, text, bold=False, italic=False,
            size=11, color=None, font="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_section_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    run.add_break()
    # horizontal rule via paragraph border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E5E7EB")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════════════════
# COVER HEADER
# ═══════════════════════════════════════════════════════════════════════════════
header = doc.add_paragraph()
shade_paragraph(header, "0F172A")
header.paragraph_format.space_before = Pt(0)
header.paragraph_format.space_after  = Pt(0)
header.paragraph_format.left_indent  = Cm(0)
add_run(header, "PRO HR  //  Astral Cassini",
        bold=True, size=22, color=WHITE, font="Calibri")

sub = doc.add_paragraph()
shade_paragraph(sub, "0F172A")
sub.paragraph_format.space_before = Pt(0)
sub.paragraph_format.space_after  = Pt(12)
add_run(sub, "Interview Preparation Guide  ·  Project Submission",
        italic=True, size=12, color=GOLD, font="Calibri")

doc.add_paragraph()  # breathing room

# ── Tip box ──────────────────────────────────────────────────────────────────
tip = doc.add_paragraph()
shade_paragraph(tip, "EEF2FF")
tip.paragraph_format.left_indent  = Cm(0.3)
tip.paragraph_format.space_before = Pt(4)
tip.paragraph_format.space_after  = Pt(8)
add_run(tip, "HOW TO USE:  ", bold=True, size=10, color=ACCENT)
add_run(tip,
        "Read each question aloud and answer from memory. "
        "Aim for fluency, not word-for-word recitation. "
        "Focus on the WHY behind every technical decision.",
        size=10, color=LIGHT_TXT)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION HELPER
# ═══════════════════════════════════════════════════════════════════════════════
def section(emoji: str, number: str, title: str):
    add_section_divider()
    p = doc.add_paragraph()
    shade_paragraph(p, "F8FAFC")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    add_run(p, f"{emoji}  {number}. ", bold=True, size=14, color=ACCENT)
    add_run(p, title,                   bold=True, size=14, color=DARK_BG)


def qa(question: str, *answer_lines):
    q = doc.add_paragraph()
    q.paragraph_format.space_before = Pt(10)
    q.paragraph_format.space_after  = Pt(2)
    add_run(q, "Q:  ", bold=True, size=11, color=ACCENT)
    add_run(q, question, bold=True, size=11, color=DARK_BG)

    for line in answer_lines:
        a = doc.add_paragraph()
        a.paragraph_format.left_indent  = Cm(0.5)
        a.paragraph_format.space_before = Pt(0)
        a.paragraph_format.space_after  = Pt(2)
        add_run(a, line, size=11, color=LIGHT_TXT)


def bullet(text: str, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(1)
    add_run(p, text, size=10.5, color=LIGHT_TXT)


def note(text: str):
    p = doc.add_paragraph()
    shade_paragraph(p, "FFFBEB")
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, "NOTE  ", bold=True, size=10, color=GOLD)
    add_run(p, text, size=10, italic=True, color=LIGHT_TXT)


# ═══════════════════════════════════════════════════════════════════════════════
# ELEVATOR PITCH
# ═══════════════════════════════════════════════════════════════════════════════
add_section_divider()
ep_label = doc.add_paragraph()
shade_paragraph(ep_label, "0F172A")
ep_label.paragraph_format.space_before = Pt(6)
ep_label.paragraph_format.space_after  = Pt(0)
add_run(ep_label, "⚡  ELEVATOR PITCH  (Deliver in first 30 seconds)",
        bold=True, size=13, color=GOLD)

ep = doc.add_paragraph()
shade_paragraph(ep, "1E293B")
ep.paragraph_format.left_indent  = Cm(0.3)
ep.paragraph_format.space_before = Pt(0)
ep.paragraph_format.space_after  = Pt(10)
add_run(ep,
        '"PRO HR is an end-to-end, AI-native recruitment automation platform. '
        'It orchestrates a multi-agent pipeline — from drafting job descriptions '
        'and sourcing candidates, all the way to conducting automated interviews '
        'and generating final hire / no-hire decisions. It is built on a LangGraph '
        'state machine for precise human-in-the-loop control, a FastAPI backend '
        'with real-time WebSocket streaming, and a Next.js dashboard. '
        'The entire system passes 150 automated tests and compiles to a production '
        'build in under 25 seconds."',
        italic=True, size=11, color=WHITE)


# ═══════════════════════════════════════════════════════════════════════════════
# 1. ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
section("🏛️", "1", "Architecture & Tech Stack")

qa("Walk me through the overall architecture.",
   "Three layers:",
   "  • Frontend — Next.js 16 App Router + custom useWebSocket hook for real-time updates.",
   "  • Backend  — FastAPI on Uvicorn, async, SQLAlchemy ORM over SQLite.",
   "  • Intelligence — LangGraph state machine orchestrating 10+ specialised AI agents.")

qa("Why FastAPI over Django or Express / Node.js?",
   "The entire AI ecosystem (LangChain, LangGraph, PyMuPDF) is Python-native — no "
   "cross-language bridge needed. FastAPI's asyncio-first design means long-running "
   "LLM calls (3–10 s) never block other HTTP requests. Django's synchronous ORM "
   "would bottleneck concurrent agent execution.")

qa("Why SQLite? Isn't that a toy database?",
   "Deliberate choice for portability and zero-config setup. All DB access goes "
   "through SQLAlchemy ORM, so the connection string in .env is the only thing "
   "that changes when promoting to PostgreSQL. Alembic migrations keep the schema "
   "versioned and reproducible.")

qa("Why Next.js with Turbopack instead of Vite?",
   "App Router lets us mix Server Components (fast initial HTML) with Client "
   "Components (live WebSocket dashboard). Turbopack cuts hot-reload time to "
   "< 200 ms. Production build compiles 15 routes and runs full TypeScript "
   "checking in under 25 seconds.")


# ═══════════════════════════════════════════════════════════════════════════════
# 2. AI & AGENTS
# ═══════════════════════════════════════════════════════════════════════════════
section("🤖", "2", "Artificial Intelligence & Agents")

q = doc.add_paragraph()
q.paragraph_format.space_before = Pt(10)
q.paragraph_format.space_after  = Pt(2)
add_run(q, "Q:  ", bold=True, size=11, color=ACCENT)
add_run(q, "Walk me through the full pipeline.", bold=True, size=11, color=DARK_BG)

stages = [
    "Intake          — HR creates a job requisition in the UI.",
    "JD Drafting     — Architect agent drafts a structured job description.",
    "JD Review       — HITL breakpoint; HR manager approves / rejects.",
    "Sourcing        — Scout agent runs BM25 + semantic search across resumes.",
    "Screening       — Screener scores candidates on skills, experience, education, culture.",
    "Shortlist Review— HITL breakpoint; human approves the shortlist.",
    "Outreach        — Personalised emails sent via SendGrid; response intent classified.",
    "Interviewing    — Coordinator schedules via Google Calendar; Interviewer conducts & scores.",
    "Decision & Offer— Decider generates hire / no-hire recs; approved candidates get offer letter.",
]
for i, s in enumerate(stages, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    add_run(p, f"{i}.  ", bold=True, size=11, color=ACCENT)
    add_run(p, s, size=11, color=LIGHT_TXT)

qa("How do you prevent hallucinations?",
   "1. Structured Outputs — every LLM call uses .with_structured_output(PydanticModel); "
   "the LLM returns JSON matching an exact schema, not free-form text.",
   "2. Grounding — Screener receives the actual extracted resume text (PyMuPDF) "
   "and the actual job requirements; scores use a deterministic weighted formula.",
   "3. Critic Loop — a separate Critic agent reviews JD drafts; if the score is "
   "below the threshold the pipeline loops back to the Architect automatically.")

qa("Why LangGraph over CrewAI or sequential chaining?",
   "• Stateful Persistence — SharedState serialised to the DB; pipeline can sleep "
   "for days awaiting approval and resume with zero data loss.",
   "• Conditional Branching — the Critic-rejection cycle is a graph loop, not "
   "expressible in a linear chain.",
   "• Native HITL breakpoints — 'interrupt before node' stops the pipeline at the "
   "three approval gates without any workarounds.")

qa("How does resume parsing and matching work?",
   "PDF → PyMuPDF extracts raw text in milliseconds.",
   "Phase 1 — BM25 Retrieval: Rank-BM25 keyword match selects top 10 candidates.",
   "Phase 2 — LLM Re-ranking: Reranker agent uses gpt-4o-mini concurrently "
   "(asyncio.gather, semaphore=5) to return matching_skills, missing_skills, match_reason.")

qa("What is the Skill Synonyms system?",
   "skill_synonyms.py maps equivalent terms: 'ML' → 'Machine Learning', "
   "'JS' → 'JavaScript', 'k8s' → 'Kubernetes'. Prevents false negatives "
   "where qualified candidates are penalised for using abbreviations.")


# ═══════════════════════════════════════════════════════════════════════════════
# 3. WEBSOCKETS
# ═══════════════════════════════════════════════════════════════════════════════
section("⚡", "3", "Real-Time WebSocket Integration")

qa("How does the dashboard update live?",
   "FastAPI manages WebSocket connections per job ID. Each LangGraph node "
   "writes to the DB and broadcasts a JSON heartbeat. The Next.js frontend's "
   "custom useWebSocket hook calls silentRefetch — pulling fresh data without "
   "disruptive UI toasts.")

qa("What was the hardest bug you fixed?",
   "The WebSocket infinite render loop on the Dashboard.",
   "Cause: WebSocket event → refetch() → toast.success() → React state change "
   "→ re-render → event still in array → refetch() again. Infinite loop.",
   "Fix: memoised silentRefetch (useCallback) that reloads data without triggering "
   "toasts. WebSocket events call silentRefetch(); the manual refresh button "
   "still calls refetch().")

qa("How do you secure WebSocket connections without HTTP headers?",
   "Short-lived ticket system: frontend calls GET /api/auth/ws-ticket → receives "
   "a signed JWT expiring in 60 s → passes it as a URL query parameter when "
   "upgrading to WebSocket. Backend validates the token on the handshake.")


# ═══════════════════════════════════════════════════════════════════════════════
# 4. SECURITY
# ═══════════════════════════════════════════════════════════════════════════════
section("🔒", "4", "Security & Error Handling")

qa("How is authentication handled?",
   "JWTs signed with python-jose (HS256). Passwords hashed with bcrypt via passlib. "
   "Four roles: admin, hr_manager, business_lead, viewer. Every route checks the "
   "decoded JWT. AUTH_DISABLED=true in .env bypasses auth for demo without "
   "touching application code.")

qa("What if an OpenAI API call fails mid-pipeline?",
   "Graceful failure at the node level. current_stage is persisted before any LLM "
   "call; an admin can retry the specific stage without losing prior work. "
   "MOCK_AGENTS=true replaces all LLM calls with deterministic responses — "
   "this is how all 150 tests pass without API keys.")

qa("What happens to candidate data? Is it private?",
   "All data is stored locally in SQLite on the server. External recipients: "
   "(1) OpenAI API under their standard DPA, (2) SendGrid for email delivery. "
   "In production: GDPR deletion endpoints, data retention policies, "
   "PII fields encrypted at rest using column-level encryption.")


# ═══════════════════════════════════════════════════════════════════════════════
# 5. ETHICS
# ═══════════════════════════════════════════════════════════════════════════════
section("⚖️", "5", "Bias, Fairness & Ethics")

qa("How does the AI avoid discrimination?",
   "1. Transcript Anonymisation — candidate names stripped; gendered pronouns "
   "(he/she/him/her) replaced with 'they/them' via regex before LLM evaluation.",
   "2. Structured Scoring — bias-prone attributes (name, age, photo) never enter "
   "the scoring input; only skills, experience years, and education.",
   "3. JD Bias Audit — Architect agent flags gendered or exclusionary language "
   "in the draft before it is published.")

qa("What if the AI makes a wrong hiring decision?",
   "The AI never makes the FINAL decision — a human always does. The system "
   "produces hire / no_hire / maybe RECOMMENDATIONS with confidence scores and "
   "reasoning. Three HITL checkpoints require active manager approval. The "
   "immutable Audit Trail logs every agent action for compliance review.")

qa("What are the ethical risks?",
   "1. Feedback Loops — historical data may encode past bias. Mitigation: audit "
   "match scores across demographic groups regularly.",
   "2. Over-reliance — managers may rubber-stamp AI recommendations. Mitigation: "
   "UI surfaces the AI's reasoning and uncertainty, not just the verdict.",
   "3. Data Privacy — resumes contain PII. Mitigation: strict retention policies, "
   "RBAC, and encrypted storage in production.")


# ═══════════════════════════════════════════════════════════════════════════════
# 6. SCALABILITY
# ═══════════════════════════════════════════════════════════════════════════════
section("📈", "6", "Scalability & System Design")

qa("How would you scale to 10,000 applicants per day?",
   "1. Database — swap SQLite for PostgreSQL (AWS RDS) via one SQLAlchemy string "
   "change. Add read replicas for the analytics dashboard.",
   "2. Task Queue — offload LangGraph execution to Celery + Redis, decoupling "
   "pipeline processing from API response times.",
   "3. Vector DB — replace BM25 with Pinecone or pgvector for sub-millisecond "
   "semantic search across millions of resumes.",
   "4. WebSockets — Redis Pub/Sub broadcasts events across multiple FastAPI "
   "instances behind a load balancer.",
   "5. LLM Cost — fine-tune a smaller open-source model (Llama 3) on HR data "
   "to cut OpenAI API costs by 80%+.")

qa("How would you add a mobile app?",
   "The FastAPI backend is fully RESTful and decoupled from Next.js. A React Native "
   "or Flutter app could consume the same endpoints. The only addition needed: "
   "Firebase Cloud Messaging for push notifications to replace browser toasts.")


# ═══════════════════════════════════════════════════════════════════════════════
# 7. DEMO WALKTHROUGH
# ═══════════════════════════════════════════════════════════════════════════════
section("🎬", "7", "Live Demo Walkthrough Script")

note("If asked to share your screen, follow this exact sequence.")

steps = [
    ("Start servers",    "cd backend && python -m uvicorn app.main:app --port 8000  |  cd frontend && npm run dev"),
    ("Open",             "http://localhost:3000"),
    ("Dashboard tab",   '"Here you can see all active pipelines, candidate counts, and live agent health."'),
    ("Kanban tab",      '"A visual Kanban board showing every job at its current pipeline stage in real time."'),
    ("Jobs → New Job",  'Create a job titled "Senior AI Engineer" in Engineering.'),
    ("Audit Trail tab", '"Every AI agent action is logged here — immutable, timestamped, and filterable."'),
    ("Candidates tab",  '"All scored candidates with match percentages, strengths, and skill gaps."'),
    ("Decisions tab",   '"Final hire / no-hire recommendations with confidence weightings and full reasoning."'),
    ("Export CSV",      '"Everything is exportable for HR compliance records."'),
]
for label, detail in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(1)
    add_run(p, f"{label}:  ", bold=True, size=11, color=ACCENT)
    add_run(p, detail, size=11, color=LIGHT_TXT)


# ═══════════════════════════════════════════════════════════════════════════════
# 8. METRICS TABLE
# ═══════════════════════════════════════════════════════════════════════════════
section("📊", "8", "Key Metrics to Quote")

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Metric"
hdr[1].text = "Value"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(11)
    shade_cell = OxmlElement("w:shd")
    shade_cell.set(qn("w:val"),   "clear")
    shade_cell.set(qn("w:color"), "auto")
    shade_cell.set(qn("w:fill"),  "0F172A")
    cell._tc.get_or_add_tcPr().append(shade_cell)

rows = [
    ("Total automated tests",           "150  (all passing, 0 failures)"),
    ("Test suite execution time",        "~11.7 seconds"),
    ("Frontend production compile time", "~11 seconds  (Turbopack)"),
    ("TypeScript check time",            "~13.7 seconds"),
    ("Static routes pre-rendered",       "15 routes"),
    ("Pipeline stages / agents",         "9 stages, 10+ specialised AI agents"),
    ("Frontend lint errors",             "0 errors  (5 minor warnings)"),
    ("Concurrent re-ranking",            "asyncio.gather, semaphore = 5"),
    ("WebSocket latency (local)",        "< 50 ms"),
    ("Skill synonym mappings",           "30+ tech term pairs"),
]
for metric, value in rows:
    row = table.add_row().cells
    row[0].text = metric
    row[1].text = value
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10.5)


# ═══════════════════════════════════════════════════════════════════════════════
# 9. BEHAVIOURAL
# ═══════════════════════════════════════════════════════════════════════════════
section("🤔", "9", "Behavioural Questions")

qa("What would you improve if you had more time?",
   '"Three things: (1) Replace BM25 with a proper vector database for semantic '
   'resume search. (2) Fine-tune a smaller open-source LLM on HR data to reduce '
   'OpenAI API dependency. (3) Add a candidate-facing portal so applicants can '
   'track their own application status in real time."')

qa("What was your biggest challenge?",
   '"Human-in-the-Loop state persistence in LangGraph. The pipeline needs to pause '
   'for days waiting for manager approval and resume exactly where it stopped — '
   'with all candidates, scores, and transcripts intact. Serialising the entire '
   'SharedState and hydrating it back without losing graph structure required '
   'very careful schema design."')

qa("How did you divide responsibilities between agents?",
   '"We followed the Single Responsibility Principle. Each agent does exactly one '
   'thing: Architect drafts, Critic reviews, Scout searches, Screener scores. '
   'This makes each agent independently testable and replaceable. Swapping OpenAI '
   'for Anthropic on just the screening step means changing one file."')

qa("Did you work alone or in a team?",
   '"I built this project with the assistance of an AI coding assistant for '
   'boilerplate and debugging. However, all architectural decisions, agent design, '
   'pipeline logic, and system design were directed and owned by me."')


# ═══════════════════════════════════════════════════════════════════════════════
# CLOSING NOTE
# ═══════════════════════════════════════════════════════════════════════════════
add_section_divider()
closing = doc.add_paragraph()
shade_paragraph(closing, "0F172A")
closing.paragraph_format.space_before = Pt(8)
closing.paragraph_format.space_after  = Pt(8)
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(closing,
        "You built something genuinely impressive. Own every decision with confidence.",
        bold=True, size=12, color=GOLD)


# ── Save ──────────────────────────────────────────────────────────────────────
out = r"c:\Users\ravip\workspace\astral-cassini\INTERVIEW_PREP.docx"
doc.save(out)
print(f"Saved -> {out}")
